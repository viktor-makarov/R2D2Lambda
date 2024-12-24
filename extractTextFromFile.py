import http.client
import urllib.parse
from io import BytesIO
from docx import Document
import os
import json
import pandas as pd
import PyPDF2
from google.api_core.client_options import ClientOptions
from google.cloud import documentai  # type: ignore
from google.auth import default
from google.oauth2 import service_account

def extracttext(docx_stream):
    doc = Document(docx_stream)
    fullText = []

    for block in doc.element.body.iterchildren():
        tag = block.tag
        if block.tag.endswith('p'):
            print('Paragraph') 
            fullText.append(block.text)
        elif block.tag.endswith('tbl'):
            # Process the table and convert to a string
            table_text = []
            print(dir(block),block.text)
            for text in block.items():
                print("text",text)

                table_text.append('\t'+text)  # Using tab as delimiter
            fullText.append('\n'.join(table_text))

    return '\n'.join(fullText)

def download_file(url):
    # Parse the URL to extract components
    url_parts = urllib.parse.urlparse(url)
    conn = http.client.HTTPSConnection(url_parts.netloc)
    conn.request("GET", url_parts.path)
    response = conn.getresponse()
    
    if response.status != 200:
        raise Exception(f"Failed to download file: {url} {response.status} {response.reason}")

    binaryObject = BytesIO(response.read())
    binaryObject.name = url_parts.path.split('/')[-1]
    
    # Read the response data into a BytesIO object
    return binaryObject

def extracttextfromtxt(txt_stream):
    return txt_stream.read().decode('utf-8')

def extracttextfromxlsx(xlsx_stream):
    df = pd.read_excel(xlsx_stream)
    dftext = df.to_string(index=False)
    
    return dftext

def extracttextfrompdf(pdf_stream):

    text = ''
    pdf_reader = PyPDF2.PdfReader(pdf_stream)
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extracttextfromtxt(filestream):
    return filestream.read().decode('utf-8')

def extracttextfromdocx(docx_stream):

    #result = extracttext(docx_stream)
    #print("print",result)

    doc = Document(docx_stream)
    
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

    doc = Document(docx_stream)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def ocrDoc(filestream,mime_type,project_id,location,processor_id):
    
    serviceaccountinfo = json.loads(os.environ['GCP_SERVICE_ACCOUNT_KEY'])

    credentials = service_account.Credentials.from_service_account_info(
        serviceaccountinfo
    )
    
    opts = ClientOptions(api_endpoint=f"{location}-documentai.googleapis.com")
    client = documentai.DocumentProcessorServiceClient(client_options=opts,credentials=credentials)
    parent = client.common_location_path(project_id, location)
    
    raw_document = documentai.RawDocument(
        content=filestream.getvalue(),
        mime_type=mime_type,  # Refer to https://cloud.google.com/document-ai/docs/file-types for supported file types
    )

    name = f"projects/{project_id}/locations/{location}/processors/{processor_id}"
    request = documentai.ProcessRequest(name=name, raw_document=raw_document)
    result = client.process_document(request=request)
    document = result.document

    return document.text

def extracttextfromfile(filestream, mime_type,project_id,location,processor_id):

    if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        text = extracttextfromdocx(filestream)
    elif mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
        text = extracttextfromxlsx(filestream)
    elif mime_type.startswith('text/') or mime_type == 'application/json':
        text = extracttextfromtxt(filestream)
    elif mime_type == 'application/pdf':
        text = extracttextfrompdf(filestream)
        noTextFound = len(text) < 10
        if noTextFound:
            text = ocrDoc(filestream,mime_type,project_id,location,processor_id)
    elif mime_type == 'image/jpeg':
        text = ocrDoc(filestream,mime_type,project_id,location,processor_id)
    else:
        raise ValueError("Unsupported file type")
    
    return text

def extractTextFromFileRouter(event, context):
  
    url = event["file_url"]
    project_id = event["project_id"]
    location = event["location"]
    processor_id = event["processor_id"]
    
    try:
        filestream = download_file(url)
    except Exception as e:
        return {
            'statusCode': 500,
            'body': {"error":'Could not fetch the file by url provided',"success":0}
        }
    mime_type  = event["file_mime_type"]

    try:
        text = extracttextfromfile(filestream, mime_type,project_id,location,processor_id)
    except Exception as e:
        return {
            'statusCode': 500,
            'body': {"error":e,"success":0}
        }

    return {
        'statusCode': 200,
        'body': {"text":text,"success":1}
    }
    	
def main(event, context):
    result = extractTextFromFileRouter(event,"")

    print("result",result)
    return result
 
#event = {"file_url":"https://r2d2storagedev.s3.eu-north-1.amazonaws.com/incoming_files/5248593849_1893.pdf","file_mime_type":"application/pdf"}

#event = {"file_url": "https://r2d2storagedev.s3.eu-north-1.amazonaws.com/incoming_files/5248593849_1866.txt", "file_mime_type": "text/plain"}

#event = {"file_url":"https://r2d2storagedev.s3.eu-north-1.amazonaws.com/incoming_files/test_xlsx.xlsx","file_mime_type":"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"}

#event = {"file_url": "https://r2d2storagedev.s3.eu-north-1.amazonaws.com/incoming_files/5248593849_1867.json", "file_mime_type": "application/json"}

#event = {"file_url":"https://r2d2storagedev.s3.eu-north-1.amazonaws.com/incoming_files/test_doc.docx","file_mime_type":"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}

event = {"file_url":"https://r2d2storagedev.s3.eu-north-1.amazonaws.com/incoming_files/test_page.jpg",
"file_mime_type":"image/jpeg",
"project_id":"119875969116",
"location": "eu",
"processor_id": "61057af8687a0f68"
}

 
if __name__ == "__main__":
    main(event, '')